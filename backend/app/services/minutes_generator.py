"""Meeting Minutes Document Generator.

Generates polished meeting minutes from analysis data including:
- Meeting details (title, date, time, duration, location)
- Attendees with roles
- Executive summary
- Agenda items with coverage status
- Key decisions
- Discussion topics with summaries
- Action items with owners, deadlines, and priorities
- Next steps

Supports output as:
- JSON (for frontend rendering)
- Word document (.docx) via python-docx
- PDF via WeasyPrint (HTML-to-PDF)
"""

import io
import uuid
import logging
from typing import List, Optional
from datetime import datetime, timezone

from sqlalchemy import select
from sqlalchemy.orm import selectinload
from sqlalchemy.ext.asyncio import AsyncSession

from app.models.meeting import Meeting, AgendaItem, MeetingAttendee
from app.models.analysis import MeetingSummary, ActionItem
from app.models.transcript import Transcript

logger = logging.getLogger(__name__)


class MinutesSection:
    """A section in the meeting minutes."""
    def __init__(self, title: str, content: str = "", items: list = None, table: list = None):
        self.title = title
        self.content = content
        self.items = items or []
        self.table = table  # list of dicts for tabular data

    def to_dict(self):
        d = {"title": self.title, "content": self.content, "items": self.items}
        if self.table:
            d["table"] = self.table
        return d


class MinutesDocument:
    """A complete meeting minutes document."""
    def __init__(self, meeting_id: uuid.UUID, title: str):
        self.meeting_id = meeting_id
        self.title = title
        self.generated_at = datetime.now(timezone.utc).isoformat()
        self.sections: List[MinutesSection] = []
        self.metadata = {}

    def add_section(self, section: MinutesSection):
        self.sections.append(section)

    def to_dict(self):
        return {
            "meeting_id": str(self.meeting_id),
            "title": self.title,
            "generated_at": self.generated_at,
            "sections": [s.to_dict() for s in self.sections],
            "metadata": self.metadata,
        }


class MinutesGeneratorService:
    """Generates meeting minutes from analysis results."""

    def __init__(self, db: AsyncSession):
        self.db = db

    async def generate_minutes(
        self,
        meeting_id: uuid.UUID,
        include_transcript_excerpts: bool = False,
    ) -> MinutesDocument:
        """Generate complete meeting minutes.

        Requires that the meeting has been analyzed (summary exists).
        """
        # Load all meeting data
        meeting = await self._load_meeting(meeting_id)
        if not meeting:
            raise ValueError(f"Meeting {meeting_id} not found")

        summary = await self._load_summary(meeting_id)
        action_items = await self._load_action_items(meeting_id)
        transcript = await self._load_transcript(meeting_id)

        minutes = MinutesDocument(meeting_id, f"Meeting Minutes — {meeting.title}")

        # Section 1: Header / Meeting Details
        header_items = [f"Meeting: {meeting.title}"]
        if meeting.date:
            header_items.append(f"Date: {meeting.date.strftime('%A, %B %d, %Y')}")
        if meeting.time:
            header_items.append(f"Time: {meeting.time.strftime('%I:%M %p')}")
        if meeting.duration_minutes:
            header_items.append(f"Duration: {meeting.duration_minutes} minutes")
        if transcript and transcript.duration_seconds:
            actual_mins = transcript.duration_seconds // 60
            header_items.append(f"Actual Duration: {actual_mins} minutes")
        if meeting.meeting_link:
            header_items.append(f"Meeting Link: {meeting.meeting_link}")

        minutes.add_section(MinutesSection(
            title="Meeting Details",
            items=header_items,
        ))

        # Section 2: Attendees
        if meeting.attendees:
            attendee_items = []
            for att in sorted(meeting.attendees, key=lambda a: (a.role != 'organizer', a.name or '')):
                entry = att.name or att.email or "Unknown"
                if att.role:
                    entry += f" ({att.role})"
                if att.rsvp_status and att.rsvp_status not in ("pending", "needsAction"):
                    entry += f" — {att.rsvp_status}"
                attendee_items.append(entry)
            minutes.add_section(MinutesSection(
                title="Attendees",
                content=f"{len(meeting.attendees)} participants",
                items=attendee_items,
            ))

        # Section 3: Executive Summary
        if summary and summary.summary_text:
            minutes.add_section(MinutesSection(
                title="Executive Summary",
                content=summary.summary_text,
            ))

        # Section 4: Key Decisions
        if summary and summary.decisions_json:
            decision_items = []
            for d in summary.decisions_json:
                entry = d.get("decision", "")
                if d.get("made_by"):
                    entry += f" (decided by {d['made_by']})"
                if d.get("context"):
                    entry += f" — {d['context']}"
                decision_items.append(entry)

            minutes.add_section(MinutesSection(
                title="Key Decisions",
                content=f"{len(decision_items)} decision(s) recorded",
                items=decision_items,
            ))

        # Section 5: Discussion Topics
        if summary and summary.topics_json:
            topic_items = []
            for t in summary.topics_json:
                entry = t.get("topic", "")
                if t.get("summary"):
                    entry += f" — {t['summary']}"
                if t.get("time_spent_estimate"):
                    entry += f" [{t['time_spent_estimate']}]"
                topic_items.append(entry)

            minutes.add_section(MinutesSection(
                title="Discussion Topics",
                content=f"{len(topic_items)} topic(s) discussed",
                items=topic_items,
            ))

        # Section 6: Action Items (table format)
        confirmed_items = [ai for ai in action_items if ai.confirmed]
        pending_items = [ai for ai in action_items if not ai.confirmed and ai.status != "declined"]

        all_actionable = confirmed_items + pending_items
        if all_actionable:
            table_data = []
            for ai in all_actionable:
                row = {
                    "task": ai.task,
                    "owner": ai.owner_name or "Unassigned",
                    "deadline": ai.deadline.strftime("%b %d, %Y") if ai.deadline else "TBD",
                    "priority": ai.priority or "medium",
                    "status": "Confirmed" if ai.confirmed else "Pending review",
                }
                table_data.append(row)

            # Also as bullet items for simpler rendering
            action_bullets = []
            for ai in all_actionable:
                entry = f"{ai.task}"
                if ai.owner_name:
                    entry += f" — Owner: {ai.owner_name}"
                if ai.deadline:
                    entry += f" — Due: {ai.deadline.strftime('%b %d, %Y')}"
                entry += f" [{ai.priority or 'medium'}]"
                if ai.confirmed:
                    entry += " ✓"
                action_bullets.append(entry)

            minutes.add_section(MinutesSection(
                title="Action Items",
                content=f"{len(all_actionable)} action item(s) ({len(confirmed_items)} confirmed, {len(pending_items)} pending review)",
                items=action_bullets,
                table=table_data,
            ))

        # Section 7: Speaker Contributions
        if summary and summary.speakers_json:
            speaker_items = []
            for s in summary.speakers_json:
                entry = s.get("name", "Unknown")
                if s.get("role"):
                    entry += f" ({s['role']})"
                if s.get("contribution_summary"):
                    entry += f" — {s['contribution_summary']}"
                speaker_items.append(entry)

            minutes.add_section(MinutesSection(
                title="Speaker Contributions",
                items=speaker_items,
            ))

        # Section 8: Next Steps
        next_steps = []
        for ai in all_actionable[:5]:  # Top 5 action items as next steps
            entry = ai.task
            if ai.owner_name:
                entry += f" ({ai.owner_name})"
            if ai.deadline:
                entry += f" by {ai.deadline.strftime('%b %d')}"
            next_steps.append(entry)

        if next_steps:
            minutes.add_section(MinutesSection(
                title="Next Steps",
                items=next_steps,
            ))

        # Metadata
        minutes.metadata = {
            "llm_provider": summary.llm_provider if summary else None,
            "llm_model": summary.llm_model if summary else None,
            "transcript_segments": transcript.segment_count if transcript else 0,
            "transcript_speakers": transcript.speaker_count if transcript else 0,
            "transcript_duration_seconds": transcript.duration_seconds if transcript else None,
            "action_items_total": len(action_items),
            "action_items_confirmed": len(confirmed_items),
            "decisions_count": len(summary.decisions_json) if summary and summary.decisions_json else 0,
            "topics_count": len(summary.topics_json) if summary and summary.topics_json else 0,
        }

        return minutes

    async def generate_minutes_docx(self, meeting_id: uuid.UUID) -> bytes:
        """Generate meeting minutes as a Word document."""
        from docx import Document as DocxDocument
        from docx.shared import Pt, Inches, RGBColor, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT

        minutes = await self.generate_minutes(meeting_id)
        doc = DocxDocument()

        # Styles
        style = doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(11)

        # Title
        title = doc.add_heading(minutes.title, level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Timestamp
        ts = doc.add_paragraph()
        ts.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = ts.add_run(f"Generated: {datetime.now(timezone.utc).strftime('%B %d, %Y at %I:%M %p UTC')}")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

        doc.add_paragraph()  # Spacer

        # Sections
        for section in minutes.sections:
            doc.add_heading(section.title, level=1)

            if section.content:
                p = doc.add_paragraph(section.content)
                p.paragraph_format.space_after = Pt(6)

            # Action Items get a table
            if section.table and section.title == "Action Items":
                table = doc.add_table(rows=1, cols=5)
                table.style = 'Light Grid Accent 1'
                table.alignment = WD_TABLE_ALIGNMENT.CENTER

                # Header row
                headers = ['Task', 'Owner', 'Deadline', 'Priority', 'Status']
                for i, header in enumerate(headers):
                    cell = table.rows[0].cells[i]
                    cell.text = header
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
                            run.font.size = Pt(9)

                # Data rows
                for row_data in section.table:
                    row = table.add_row()
                    row.cells[0].text = row_data.get("task", "")
                    row.cells[1].text = row_data.get("owner", "")
                    row.cells[2].text = row_data.get("deadline", "")
                    row.cells[3].text = row_data.get("priority", "")
                    row.cells[4].text = row_data.get("status", "")
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.size = Pt(9)

                doc.add_paragraph()  # Spacer after table
            else:
                # Regular bullet items
                for item in section.items:
                    doc.add_paragraph(item, style='List Bullet')

            doc.add_paragraph()  # Spacer between sections

        # Footer
        doc.add_paragraph()
        footer = doc.add_paragraph()
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer.add_run("— End of Minutes —")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
        run.italic = True

        # Save to bytes
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

    async def generate_minutes_pdf(self, meeting_id: uuid.UUID) -> bytes:
        """Generate meeting minutes as a PDF document via HTML template."""
        minutes = await self.generate_minutes(meeting_id)

        html = self._render_html(minutes)

        try:
            from weasyprint import HTML
            pdf_bytes = HTML(string=html).write_pdf()
            return pdf_bytes
        except ImportError:
            logger.warning("WeasyPrint not available, falling back to HTML response")
            return html.encode('utf-8')
        except Exception as e:
            logger.error(f"PDF generation failed: {e}")
            raise

    def _render_html(self, minutes: MinutesDocument) -> str:
        """Render minutes as an HTML document for PDF conversion."""
        sections_html = ""
        for section in minutes.sections:
            sections_html += f'<h2>{section.title}</h2>\n'
            if section.content:
                sections_html += f'<p class="content">{section.content}</p>\n'

            if section.table and section.title == "Action Items":
                sections_html += '<table><thead><tr>'
                sections_html += '<th>Task</th><th>Owner</th><th>Deadline</th><th>Priority</th><th>Status</th>'
                sections_html += '</tr></thead><tbody>\n'
                for row in section.table:
                    pri_cls = f'priority-{row.get("priority", "medium")}'
                    sections_html += f'<tr><td>{row.get("task","")}</td><td>{row.get("owner","")}</td>'
                    sections_html += f'<td>{row.get("deadline","")}</td><td class="{pri_cls}">{row.get("priority","")}</td>'
                    sections_html += f'<td>{row.get("status","")}</td></tr>\n'
                sections_html += '</tbody></table>\n'
            elif section.items:
                sections_html += '<ul>\n'
                for item in section.items:
                    sections_html += f'  <li>{item}</li>\n'
                sections_html += '</ul>\n'

        return f"""<!DOCTYPE html>
<html><head><meta charset="utf-8">
<style>
  body {{ font-family: 'Segoe UI', Calibri, Arial, sans-serif; max-width: 800px; margin: 0 auto; padding: 40px; color: #1d1b16; line-height: 1.6; }}
  h1 {{ text-align: center; font-size: 24px; margin-bottom: 4px; }}
  .timestamp {{ text-align: center; color: #999; font-size: 11px; margin-bottom: 30px; }}
  h2 {{ font-size: 16px; border-bottom: 2px solid #E8622C; padding-bottom: 4px; margin-top: 28px; color: #1d1b16; }}
  .content {{ font-size: 13px; color: #444; white-space: pre-wrap; }}
  ul {{ padding-left: 20px; }}
  li {{ font-size: 13px; margin-bottom: 6px; color: #333; }}
  table {{ width: 100%; border-collapse: collapse; font-size: 12px; margin: 12px 0; }}
  th {{ background: #f5f4f0; padding: 8px 10px; text-align: left; font-weight: 600; border: 1px solid #e0e0e0; }}
  td {{ padding: 7px 10px; border: 1px solid #e0e0e0; vertical-align: top; }}
  .priority-high {{ color: #dc2626; font-weight: 600; }}
  .priority-medium {{ color: #b45309; }}
  .priority-low {{ color: #2d8544; }}
  .footer {{ text-align: center; color: #999; font-style: italic; font-size: 11px; margin-top: 40px; }}
</style></head><body>
<h1>{minutes.title}</h1>
<p class="timestamp">Generated: {datetime.now(timezone.utc).strftime('%B %d, %Y at %I:%M %p UTC')}</p>
{sections_html}
<p class="footer">— End of Minutes —</p>
</body></html>"""

    # ============================================
    # DATA LOADING HELPERS
    # ============================================

    async def _load_meeting(self, meeting_id: uuid.UUID) -> Optional[Meeting]:
        result = await self.db.execute(
            select(Meeting)
            .options(
                selectinload(Meeting.agenda_items),
                selectinload(Meeting.attendees),
            )
            .where(Meeting.id == meeting_id)
        )
        return result.scalar_one_or_none()

    async def _load_summary(self, meeting_id: uuid.UUID) -> Optional[MeetingSummary]:
        result = await self.db.execute(
            select(MeetingSummary).where(MeetingSummary.meeting_id == meeting_id)
        )
        return result.scalar_one_or_none()

    async def _load_action_items(self, meeting_id: uuid.UUID) -> List[ActionItem]:
        result = await self.db.execute(
            select(ActionItem)
            .where(ActionItem.meeting_id == meeting_id)
            .order_by(ActionItem.priority.desc(), ActionItem.created_at)
        )
        return list(result.scalars().all())

    async def _load_transcript(self, meeting_id: uuid.UUID) -> Optional[Transcript]:
        result = await self.db.execute(
            select(Transcript).where(Transcript.meeting_id == meeting_id)
        )
        return result.scalar_one_or_none()
