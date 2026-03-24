"""Pre-Meeting Briefing Package Generator.

Compiles approved documents, structured agenda, attendee information, and
outstanding action items from previous meetings into a consolidated briefing.
Generates output as JSON (for frontend rendering) and optionally as a Word document.
"""

import io
import uuid
import logging
from typing import List, Optional
from datetime import datetime, timezone

from sqlalchemy import select
from sqlalchemy.orm import selectinload
from sqlalchemy.ext.asyncio import AsyncSession

from app.models.meeting import Meeting, AgendaItem, MeetingAttendee, Document
from app.models.analysis import ActionItem

logger = logging.getLogger(__name__)


class BriefingSection:
    """A section in the briefing package."""
    def __init__(self, title: str, content: str = "", items: list = None):
        self.title = title
        self.content = content
        self.items = items or []

    def to_dict(self):
        return {"title": self.title, "content": self.content, "items": self.items}


class BriefingPackage:
    """A complete pre-meeting briefing package."""
    def __init__(self, meeting_id: uuid.UUID, meeting_title: str):
        self.meeting_id = meeting_id
        self.meeting_title = meeting_title
        self.generated_at = datetime.now(timezone.utc).isoformat()
        self.sections: List[BriefingSection] = []
        self.metadata = {}

    def add_section(self, section: BriefingSection):
        self.sections.append(section)

    def to_dict(self):
        return {
            "meeting_id": str(self.meeting_id),
            "meeting_title": self.meeting_title,
            "generated_at": self.generated_at,
            "sections": [s.to_dict() for s in self.sections],
            "metadata": self.metadata,
        }


class BriefingGeneratorService:
    """Generates pre-meeting briefing packages from meeting data."""

    def __init__(self, db: AsyncSession):
        self.db = db

    async def generate_briefing(
        self,
        meeting_id: uuid.UUID,
        include_outstanding_actions: bool = True,
        include_documents: bool = True,
    ) -> BriefingPackage:
        """Generate a complete briefing package for a meeting.

        Pulls together:
        1. Meeting details (title, date, time, link)
        2. Structured agenda with time allocations
        3. Attendee list with roles
        4. Approved documents linked to agenda items
        5. Outstanding action items from previous meetings (same organizer)
        """
        # Load meeting with all relationships
        result = await self.db.execute(
            select(Meeting)
            .options(
                selectinload(Meeting.agenda_items),
                selectinload(Meeting.attendees),
                selectinload(Meeting.documents),
            )
            .where(Meeting.id == meeting_id)
        )
        meeting = result.scalar_one_or_none()

        if not meeting:
            raise ValueError(f"Meeting {meeting_id} not found")

        briefing = BriefingPackage(meeting_id, meeting.title)

        # Section 1: Meeting Overview
        overview_items = []
        if meeting.date:
            overview_items.append(f"Date: {meeting.date.strftime('%A, %B %d, %Y')}")
        if meeting.time:
            overview_items.append(f"Time: {meeting.time.strftime('%I:%M %p')}")
        if meeting.duration_minutes:
            overview_items.append(f"Duration: {meeting.duration_minutes} minutes")
        if meeting.meeting_link:
            overview_items.append(f"Meeting Link: {meeting.meeting_link}")
        if meeting.notes:
            overview_items.append(f"Notes: {meeting.notes}")

        briefing.add_section(BriefingSection(
            title="Meeting Overview",
            content=meeting.title,
            items=overview_items,
        ))

        # Section 2: Agenda
        if meeting.agenda_items:
            agenda_sorted = sorted(meeting.agenda_items, key=lambda a: a.item_order)
            agenda_items = []
            for item in agenda_sorted:
                entry = f"{item.item_order + 1}. {item.title}"
                if item.time_allocation_minutes:
                    entry += f" ({item.time_allocation_minutes} min)"
                if item.description:
                    entry += f" — {item.description}"
                agenda_items.append(entry)

            briefing.add_section(BriefingSection(
                title="Agenda",
                content=f"{len(agenda_sorted)} agenda items",
                items=agenda_items,
            ))

        # Section 3: Attendees
        if meeting.attendees:
            attendee_items = []
            for att in meeting.attendees:
                entry = att.name or att.email or "Unknown"
                if att.role:
                    entry += f" ({att.role})"
                if att.email:
                    entry += f" — {att.email}"
                if att.rsvp_status and att.rsvp_status != "pending":
                    entry += f" [{att.rsvp_status}]"
                attendee_items.append(entry)

            briefing.add_section(BriefingSection(
                title="Attendees",
                content=f"{len(meeting.attendees)} participants",
                items=attendee_items,
            ))

        # Section 4: Approved Documents
        if include_documents and meeting.documents:
            approved_docs = [d for d in meeting.documents if d.approved]
            if approved_docs:
                doc_items = []
                for doc in approved_docs:
                    entry = doc.file_name
                    if doc.file_url:
                        entry += f" — {doc.file_url}"
                    if doc.agenda_item:
                        entry += f" (linked to: {doc.agenda_item.title})"
                    doc_items.append(entry)

                briefing.add_section(BriefingSection(
                    title="Reference Documents",
                    content=f"{len(approved_docs)} approved documents",
                    items=doc_items,
                ))

        # Section 5: Outstanding Action Items
        if include_outstanding_actions:
            outstanding = await self._get_outstanding_actions(meeting_id)
            if outstanding:
                action_items = []
                for ai in outstanding:
                    entry = ai.task
                    if ai.owner_name:
                        entry += f" — Owner: {ai.owner_name}"
                    if ai.deadline:
                        entry += f" — Due: {ai.deadline.strftime('%b %d, %Y')}"
                    if ai.status and ai.status != "pending":
                        entry += f" [{ai.status}]"
                    action_items.append(entry)

                briefing.add_section(BriefingSection(
                    title="Outstanding Action Items",
                    content=f"{len(outstanding)} items from previous meetings",
                    items=action_items,
                ))

        # Metadata
        briefing.metadata = {
            "agenda_items_count": len(meeting.agenda_items) if meeting.agenda_items else 0,
            "attendees_count": len(meeting.attendees) if meeting.attendees else 0,
            "documents_count": len([d for d in meeting.documents if d.approved]) if meeting.documents else 0,
            "outstanding_actions_count": len(await self._get_outstanding_actions(meeting_id)) if include_outstanding_actions else 0,
        }

        logger.info(f"Generated briefing for meeting {meeting_id}: {len(briefing.sections)} sections")
        return briefing

    async def generate_briefing_docx(
        self, meeting_id: uuid.UUID
    ) -> bytes:
        """Generate briefing as a Word document.

        Returns the .docx file as bytes.
        """
        from docx import Document as DocxDocument
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        briefing = await self.generate_briefing(meeting_id)
        doc = DocxDocument()

        # Title
        title = doc.add_heading(briefing.meeting_title, level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Subtitle
        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle.add_run("Pre-Meeting Briefing Package")
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0x6B, 0x68, 0x60)

        generated = doc.add_paragraph()
        generated.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = generated.add_run(f"Generated: {datetime.now(timezone.utc).strftime('%B %d, %Y at %I:%M %p UTC')}")
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x9C, 0x9A, 0x93)

        doc.add_paragraph()  # Spacer

        # Sections
        for section in briefing.sections:
            doc.add_heading(section.title, level=1)

            if section.content:
                doc.add_paragraph(section.content)

            for item in section.items:
                p = doc.add_paragraph(item, style="List Bullet")

            doc.add_paragraph()  # Spacer between sections

        # Save to bytes
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

    async def _get_outstanding_actions(
        self, meeting_id: uuid.UUID
    ) -> List[ActionItem]:
        """Get all incomplete action items from previous meetings.

        Returns pending/in-progress action items that haven't been completed,
        which may be relevant for the current meeting.
        """
        # Get the current meeting to find organizer context
        result = await self.db.execute(
            select(Meeting).where(Meeting.id == meeting_id)
        )
        meeting = result.scalar_one_or_none()
        if not meeting:
            return []

        # Find all incomplete action items from other meetings
        result = await self.db.execute(
            select(ActionItem)
            .where(
                ActionItem.meeting_id != meeting_id,
                ActionItem.status.in_(["pending", "in_progress"]),
                ActionItem.confirmed == True,
            )
            .order_by(ActionItem.deadline.asc().nullslast(), ActionItem.created_at.desc())
            .limit(20)
        )
        return list(result.scalars().all())
